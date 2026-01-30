from __future__ import annotations

from dataclasses import dataclass
from typing import Iterable
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document


LIST_URL = "https://news.rthk.hk/rthk/ch/latest-news/world-news.htm"
LIST_AJAX_URL = (
    "https://news.rthk.hk/rthk/webpageCache/services/loadModNewsShowSp2List.php"
)
LIST_AJAX_PARAMS = {
    "lang": "zh-TW",
    "cat": "4",  # 国际
    "newsCount": 60,
    "dayShiftMode": "1",
    "archive_date": "",
}
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}


@dataclass
class ArticleLink:
    title: str
    url: str
    time_text: str | None = None


def fetch_html(url: str, params: dict | None = None) -> str:
    resp = requests.get(url, headers=HEADERS, params=params, timeout=20)
    resp.raise_for_status()
    # Let requests decode using detected encoding to avoid mojibake.
    resp.encoding = resp.apparent_encoding or resp.encoding
    return resp.text


def normalize_url(href: str) -> str:
    return urljoin(LIST_URL, href)


def parse_list(html: str) -> list[ArticleLink]:
    soup = BeautifulSoup(html, "html.parser")
    results: list[ArticleLink] = []

    top_container = soup.select_one(".catTopNewsContainer")
    if top_container:
        top_link = top_container.select_one("a[href]")
        top_title = top_container.select_one(".catTopNewsTitleText a")
        top_time = top_container.select_one(".catTopTime")
        if top_link and top_title:
            results.append(
                ArticleLink(
                    title=top_title.get_text(strip=True),
                    url=normalize_url(top_link["href"]),
                    time_text=top_time.get_text(strip=True) if top_time else None,
                )
            )

    for link in soup.select(".ns2-title a[href]"):
        title = link.get_text(strip=True)
        url = normalize_url(link["href"])
        time_text = None
        inner = link.find_parent(class_="ns2-inner")
        if inner:
            time_el = inner.select_one(".ns2-created")
            if time_el:
                time_text = time_el.get_text(strip=True)
        results.append(ArticleLink(title=title, url=url, time_text=time_text))

    return dedupe_by_url(results)


def dedupe_by_url(items: Iterable[ArticleLink]) -> list[ArticleLink]:
    seen: set[str] = set()
    results: list[ArticleLink] = []
    for item in items:
        if item.url in seen:
            continue
        seen.add(item.url)
        results.append(item)
    return results


def parse_detail(html: str) -> dict[str, str | None]:
    soup = BeautifulSoup(html, "html.parser")
    title_el = soup.select_one("h2.itemTitle")
    body_el = soup.select_one(".itemFullText")
    title = title_el.get_text(strip=True) if title_el else None

    body = None
    if body_el:
        for br in body_el.find_all("br"):
            br.replace_with("\n")
        raw = body_el.get_text("\n")
        body = "\n".join(line.strip() for line in raw.splitlines() if line.strip())

    return {"title": title, "body": body}


def build_docx(items: Iterable[ArticleLink], output_path: str) -> None:
    doc = Document()
    for item in items:
        detail_html = fetch_html(item.url)
        detail = parse_detail(detail_html)
        title = detail["title"] or item.title
        body = detail["body"] or ""

        doc.add_heading(title, level=1)
        if body:
            for paragraph in body.splitlines():
                doc.add_paragraph(paragraph)
        else:
            doc.add_paragraph("(正文为空)")
        doc.add_paragraph("")

    doc.save(output_path)


def main() -> None:
    list_html = fetch_html(LIST_AJAX_URL, params=LIST_AJAX_PARAMS)
    links = parse_list(list_html)
    top10 = links[:10]
    output_path = "rthk_world_top10.docx"
    build_docx(top10, output_path)
    print(f"Parsed {len(links)} article(s). Docx saved to {output_path}.")


if __name__ == "__main__":
    main()
